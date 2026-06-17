"""Load documents from DATASET_DIR, attach classification labels, and chunk them.

Supported: .docx, .pdf, .csv, .xlsx, .txt, .md. Files not found in the
CLASSIFICATION map fall back to DEFAULT_LEVEL.
"""
import os
import csv
import config


def _read_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(path):
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _read_xlsx(path, max_rows=300):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i > max_rows:
                out.append("... (truncated)")
                break
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _read_csv(path, max_rows=400):
    out = []
    with open(path, encoding="utf-8", errors="replace") as f:
        for i, row in enumerate(csv.reader(f)):
            if i > max_rows:
                out.append("... (truncated)")
                break
            out.append(", ".join(row))
    return "\n".join(out)


def _read_text(path):
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


_READERS = {
    ".docx": _read_docx, ".pdf": _read_pdf, ".xlsx": _read_xlsx,
    ".csv": _read_csv, ".txt": _read_text, ".md": _read_text,
}


def classify(basename):
    """Return (level, allowed_roles) for a document basename."""
    if basename in config.CLASSIFICATION:
        return config.CLASSIFICATION[basename]
    return (config.DEFAULT_LEVEL, None)


def _chunk(text, size=900, overlap=150):
    text = text.strip()
    if not text:
        return []
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
    return chunks


def load_documents(dataset_dir=None):
    """Walk the dataset and return a list of chunk dicts:
    {id, doc, level, level_name, allowed_roles, text}
    """
    dataset_dir = dataset_dir or config.DATASET_DIR
    chunks = []
    if not os.path.isdir(dataset_dir):
        return chunks
    cid = 0
    for root, _dirs, files in os.walk(dataset_dir):
        for fn in sorted(files):
            ext = os.path.splitext(fn)[1].lower()
            reader = _READERS.get(ext)
            if not reader:
                continue
            path = os.path.join(root, fn)
            try:
                text = reader(path)
            except Exception as e:  # pragma: no cover - corrupt/locked file
                print(f"[ingest] skip {fn}: {e}")
                continue
            level, roles = classify(fn)
            for piece in _chunk(text):
                chunks.append({
                    "id": cid, "doc": fn, "level": level,
                    "level_name": config.LEVEL_NAMES[level],
                    "allowed_roles": roles, "text": piece,
                })
                cid += 1
    return chunks
